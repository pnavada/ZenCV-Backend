from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any


class DocumentHandler:
    @staticmethod
    def create_docx(content: Dict[str, Any]) -> Document:
        print(content)
        doc = Document()
        
        # Document formatting
        section = doc.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

        def add_section_heading(text: str, space_before: bool = True):
            if space_before:
                doc.add_paragraph().space_after = Pt(0)
            heading = doc.add_paragraph()
            run = heading.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            heading.space_before = Pt(0)
            heading.space_after = Pt(6)
            return heading

        def add_horizontal_line():
            p = doc.add_paragraph()
            p.space_before = Pt(3)
            p.space_after = Pt(6)
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            p._element.get_or_add_pPr().append(pBdr)

        def set_right_tab(paragraph):
            # Set a right-aligned tab stop at 6.4 inches (adjusting for margins)
            paragraph_format = paragraph.paragraph_format
            tab_stops = paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT)

        # Header
        if 'contact' in content:
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(content['contact']['name'])
            name_run.bold = True
            name_run.font.size = Pt(16)
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_paragraph.space_after = Pt(4)

            contact_paragraph = doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_parts = []
            if 'email' in content['contact']:
                contact_parts.append(content['contact']['email'])
            if 'phone' in content['contact']:
                contact_parts.append(content['contact']['phone'])
            if 'location' in content['contact']:
                contact_parts.append(content['contact']['location'])
            if 'linkedin' in content['contact']:
                contact_parts.append(content['contact']['linkedin'])
            
            contact_text = ' | '.join(contact_parts)
            contact_run = contact_paragraph.add_run(contact_text)
            contact_run.font.size = Pt(10.5)
            contact_paragraph.space_after = Pt(6)

        add_horizontal_line()

        # Summary
        if 'summary' in content:
            add_section_heading('PROFESSIONAL SUMMARY', space_before=False)
            summary = doc.add_paragraph()
            summary.add_run(content['summary']).font.size = Pt(10.5)
            summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            summary.space_after = Pt(12)

        # Experience
        if 'experience' in content:
            add_section_heading('PROFESSIONAL EXPERIENCE')
            for exp in content['experience']:
                p = doc.add_paragraph()
                set_right_tab(p)
                
                # Title and Company on left
                title_company = p.add_run(f"{exp['title']} - {exp['company']}")
                title_company.bold = True
                title_company.font.size = Pt(11)
                
                # Date on right using tab
                p.add_run('\t')  # Add tab
                p.add_run(exp['dates']).font.size = Pt(10.5)
                p.space_after = Pt(4)

                # Bullets
                for point in exp['points']:
                    bullet = doc.add_paragraph(style='List Bullet')
                    bullet.add_run(point).font.size = Pt(10.5)
                    bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    bullet.space_after = Pt(2)
                
                if exp != content['experience'][-1]:
                    doc.add_paragraph().space_after = Pt(6)

        # Education
        if 'education' in content:
            add_section_heading('EDUCATION')
            for edu in content['education']:
                p = doc.add_paragraph()
                set_right_tab(p)
                
                # Degree on left
                degree = p.add_run(f"{edu['degree']}")
                degree.bold = True
                degree.font.size = Pt(11)
                
                # Dates on right using tab
                p.add_run('\t')
                p.add_run(f"{edu['dates']}").font.size = Pt(10.5)
                p.space_after = Pt(2)

                # School on next line
                school = doc.add_paragraph()
                school.add_run(edu['school']).font.size = Pt(10.5)
                school.space_after = Pt(4)

                if 'courses' in edu and edu['courses']:
                    courses = doc.add_paragraph()
                    courses.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    courses_run = courses.add_run('Relevant Coursework: ')
                    courses_run.bold = True
                    courses_run.font.size = Pt(10.5)
                    courses.add_run(', '.join(edu['courses'])).font.size = Pt(10.5)
                    courses.space_after = Pt(8)

        # Skills
        if 'skills' in content:
            add_section_heading('TECHNICAL SKILLS')
            for category, skills in content['skills'].items():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(f"{category}: ").bold = True
                p.add_run(', '.join(skills))
                for run in p.runs:
                    run.font.size = Pt(10.5)
                p.space_after = Pt(4)

        # Projects
        if 'projects' in content:
            add_section_heading('PERSONAL PROJECTS')
            for project in content['projects']:
                p = doc.add_paragraph()
                set_right_tab(p)
                
                # Project name on left
                name = p.add_run(project['name'])
                name.bold = True
                name.font.size = Pt(11)
                
                # Date on right using tab

                if 'dates' in project and project['dates'] and project['dates'] != 'None':
                    p.add_run('\t')
                    p.add_run(f"{project['dates']}").font.size = Pt(10.5)
                    p.space_after = Pt(2)

                if 'technologies' in project:
                    tech = doc.add_paragraph()
                    tech.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    tech.add_run('Technologies: ').bold = True
                    tech.add_run(', '.join(project['technologies']))
                    for run in tech.runs:
                        run.font.size = Pt(10.5)
                    tech.space_after = Pt(4)

                if 'points' in project:
                    for point in project['points']:
                        bullet = doc.add_paragraph(style='List Bullet')
                        bullet.add_run(point).font.size = Pt(10.5)
                        bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        bullet.space_after = Pt(2)

                # if project != content['projects'][-1]:
                #     doc.add_paragraph().space_after = Pt(6)

        # Volunteering
        if 'volunteering' in content:
            add_section_heading('VOLUNTEER EXPERIENCE')
            for vol in content['volunteering']:
                p = doc.add_paragraph()
                set_right_tab(p)
                
                # Role and Organization on left
                role_org = p.add_run(f"{vol['role']} - {vol['organization']}")
                role_org.bold = True
                role_org.font.size = Pt(11)
                
                # Date on right using tab
                p.add_run('\t')
                p.add_run(vol['dates']).font.size = Pt(10.5)
                p.space_after = Pt(4)

                if 'points' in vol:
                    for point in vol['points']:
                        bullet = doc.add_paragraph(style='List Bullet')
                        bullet.add_run(point).font.size = Pt(10.5)
                        bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        bullet.space_after = Pt(2)

                if vol != content['volunteering'][-1]:
                    doc.add_paragraph().space_after = Pt(6)

        # Honors & Awards
        if 'honors' in content:
            add_section_heading('HONORS & AWARDS')
            for honor in content['honors']:
                p = doc.add_paragraph()
                set_right_tab(p)
                
                # Award and Issuer on left
                title = p.add_run(f"{honor['title']} - {honor['issuer']}")
                title.bold = True
                title.font.size = Pt(11)
                
                # Date on right using tab
                if 'date' in honor and honor['date'] and honor['date'] != 'None':
                    p.add_run('\t')
                    p.add_run(f"{honor['date']}").font.size = Pt(10.5)
                    p.space_after = Pt(2)

                if 'description' in honor:
                    desc = doc.add_paragraph()
                    desc.add_run(honor['description']).font.size = Pt(10.5)
                    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    desc.space_after = Pt(4)

        return doc

# if __name__ == "__main__":
#     # Sample content for testing
#     sample_content = {
#         'contact': {
#             'name': 'John Doe',
#             'email': 'john.doe@email.com',
#             'phone': '(123) 456-7890',
#             'location': 'New York, NY',
#             'linkedin': 'linkedin.com/in/johndoe'
#         },
#         'summary': 'Experienced software engineer with expertise in Python and web development...',
#         'experience': [{
#             'title': 'Senior Software Engineer',
#             'company': 'Tech Corp',
#             'dates': 'Jan 2020 - Present',
#             'points': [
#                 'Led development of microservices architecture serving 1M+ users',
#                 'Implemented CI/CD pipeline reducing deployment time by 50%'
#             ]
#         }],
#         'education': [{
#             'degree': 'BS Computer Science',
#             'school': 'University of Technology',
#             'dates': '2015 - 2019',
#             'courses': [
#                 'Data Structures',
#                 'Algorithms',
#                 'Machine Learning',
#                 'Database Systems'
#             ]
#         }],
#         'skills': {
#             'Programming': ['Python', 'JavaScript', 'Java'],
#             'Web Technologies': ['React', 'Node.js', 'Django'],
#             'Tools & Platforms': ['AWS', 'Docker', 'Kubernetes']
#         },
#         'projects': [{
#             'name': 'AI-Powered Resume Builder',
#             'dates': '2023',
#             'description': 'Built a web application that uses AI to customize resumes',
#             'points': ['Integrated OpenAI GPT-4 for content generation'],
#             'technologies': ['Python', 'React', 'FastAPI']
#         }],
#         'volunteering': [{
#             'organization': 'Code for Good',
#             'role': 'Technical Mentor',
#             'dates': '2022 - Present',
#             'points': ['Mentored 20+ junior developers']
#         }],
#         'honors': [{
#             'title': 'Outstanding Technical Achievement',
#             'issuer': 'Tech Corp',
#             'date': '2022',
#             'description': 'Awarded for innovative solution to scaling challenges'
#         }]
#     }

#     # Create and save document
#     doc = DocumentHandler.create_docx(sample_content)
#     doc.save('sample_resume.docx')
